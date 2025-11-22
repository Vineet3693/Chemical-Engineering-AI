"""
Export Handler Module
Handles exporting responses to PDF and DOCX formats
"""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
from pathlib import Path
from utils.helpers import get_timestamp
from config.settings import settings
from utils.logger import get_logger, log_execution_time
import re

# Initialize logger
logger = get_logger(__name__)


class ExportHandler:
    """Handle exporting responses to PDF and DOCX formats"""
    
    def __init__(self):
        """Initialize export handler"""
        # Ensure export directory exists
        settings.EXPORT_DIR.mkdir(exist_ok=True)
        logger.info(f"ExportHandler initialized | Export directory: {settings.EXPORT_DIR}")
    
    def _clean_text_for_export(self, text: str) -> str:
        """
        Clean text for export (remove markdown, etc.)
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        # Remove markdown bold
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        
        # Remove markdown italic
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        
        # Remove markdown code blocks
        text = re.sub(r'`(.*?)`', r'\1', text)
        
        return text
    
    @log_execution_time
    def export_to_pdf(
        self,
        query: str,
        answer: str,
        citations: str,
        mode: str,
        filename: str = None
    ) -> str:
        """
        Export response to PDF
        
        Args:
            query: User query
            answer: Generated answer
            citations: Citation text
            mode: Response mode ('book_based' or 'general_knowledge')
            filename: Optional custom filename
            
        Returns:
            Path to generated PDF file
        """
        logger.info(f"Exporting to PDF | Mode: {mode} | Query length: {len(query)} chars")
        
        # Generate filename
        if not filename:
            timestamp = get_timestamp().replace(':', '-').replace(' ', '_')
            filename = f"response_{timestamp}.pdf"
        
        filepath = settings.EXPORT_DIR / filename
        
        try:
            # Create PDF
            doc = SimpleDocTemplate(str(filepath), pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=HexColor('#003366'),
                spaceAfter=12,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=HexColor('#0066CC'),
                spaceAfter=6,
                spaceBefore=12
            )
            
            # Title
            story.append(Paragraph("Chemical Engineering RAG System", title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Metadata
            meta_text = f"<b>Generated:</b> {get_timestamp()}<br/><b>Mode:</b> {mode.replace('_', ' ').title()}"
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(Spacer(1, 0.3*inch))
            
            # Query
            story.append(Paragraph("Question", heading_style))
            query_clean = self._clean_text_for_export(query)
            story.append(Paragraph(query_clean, styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
            
            # Answer
            story.append(Paragraph("Answer", heading_style))
            answer_clean = self._clean_text_for_export(answer)
            
            # Split answer into paragraphs
            for para in answer_clean.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Citations
            if citations:
                story.append(Spacer(1, 0.2*inch))
                story.append(Paragraph("References", heading_style))
                citations_clean = self._clean_text_for_export(citations)
                
                for citation in citations_clean.split('\n'):
                    if citation.strip():
                        story.append(Paragraph(citation.strip(), styles['Normal']))
                        story.append(Spacer(1, 0.05*inch))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF exported successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Failed to export PDF: {type(e).__name__}: {str(e)}")
            raise
    
    @log_execution_time
    def export_to_docx(
        self,
        query: str,
        answer: str,
        citations: str,
        mode: str,
        filename: str = None
    ) -> str:
        """
        Export response to DOCX
        
        Args:
            query: User query
            answer: Generated answer
            citations: Citation text
            mode: Response mode ('book_based' or 'general_knowledge')
            filename: Optional custom filename
            
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Exporting to DOCX | Mode: {mode} | Query length: {len(query)} chars")
        
        # Generate filename
        if not filename:
            timestamp = get_timestamp().replace(':', '-').replace(' ', '_')
            filename = f"response_{timestamp}.docx"
        
        filepath = settings.EXPORT_DIR / filename
        
        try:
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Chemical Engineering RAG System', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            meta = doc.add_paragraph()
            meta.add_run('Generated: ').bold = True
            meta.add_run(f"{get_timestamp()}\n")
            meta.add_run('Mode: ').bold = True
            meta.add_run(mode.replace('_', ' ').title())
            
            doc.add_paragraph()  # Spacer
            
            # Query
            doc.add_heading('Question', 1)
            query_clean = self._clean_text_for_export(query)
            doc.add_paragraph(query_clean)
            
            # Answer
            doc.add_heading('Answer', 1)
            answer_clean = self._clean_text_for_export(answer)
            
            # Split answer into paragraphs
            for para in answer_clean.split('\n\n'):
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # Citations
            if citations:
                doc.add_heading('References', 1)
                citations_clean = self._clean_text_for_export(citations)
                
                for citation in citations_clean.split('\n'):
                    if citation.strip():
                        doc.add_paragraph(citation.strip(), style='List Bullet')
            
            # Save document
            doc.save(str(filepath))
            
            logger.info(f"DOCX exported successfully: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"Failed to export DOCX: {type(e).__name__}: {str(e)}")
            raise
    
    def export_chat_history(
        self,
        chat_history: List[Dict[str, Any]],
        format: str = 'pdf',
        filename: str = None
    ) -> str:
        """
        Export entire chat history
        
        Args:
            chat_history: List of Q&A dictionaries
            format: 'pdf' or 'docx'
            filename: Optional custom filename
            
        Returns:
            Path to exported file
        """
        logger.info(f"Exporting chat history | Format: {format} | Items: {len(chat_history)}")
        
        if format == 'pdf':
            return self._export_history_pdf(chat_history, filename)
        else:
            return self._export_history_docx(chat_history, filename)
    
    def _export_history_pdf(self, chat_history: List[Dict[str, Any]], filename: str = None) -> str:
        """Export chat history to PDF"""
        if not filename:
            timestamp = get_timestamp().replace(':', '-').replace(' ', '_')
            filename = f"chat_history_{timestamp}.pdf"
        
        filepath = settings.EXPORT_DIR / filename
        
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=HexColor('#003366'),
            alignment=TA_CENTER
        )
        
        story.append(Paragraph("Chemical Engineering RAG - Chat History", title_style))
        story.append(Paragraph(f"Generated: {get_timestamp()}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Add each Q&A
        for i, item in enumerate(chat_history, 1):
            story.append(Paragraph(f"<b>Q{i}:</b> {item['query']}", styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(f"<b>A{i}:</b> {item['answer']}", styles['Normal']))
            
            if item.get('citations'):
                story.append(Spacer(1, 0.05*inch))
                story.append(Paragraph(f"<i>{item['citations']}</i>", styles['Italic']))
            
            story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        logger.info(f"Chat history PDF exported: {filepath}")
        return str(filepath)
    
    def _export_history_docx(self, chat_history: List[Dict[str, Any]], filename: str = None) -> str:
        """Export chat history to DOCX"""
        if not filename:
            timestamp = get_timestamp().replace(':', '-').replace(' ', '_')
            filename = f"chat_history_{timestamp}.docx"
        
        filepath = settings.EXPORT_DIR / filename
        
        doc = Document()
        title = doc.add_heading('Chemical Engineering RAG - Chat History', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated: {get_timestamp()}")
        
        # Add each Q&A
        for i, item in enumerate(chat_history, 1):
            doc.add_heading(f"Question {i}", 2)
            doc.add_paragraph(item['query'])
            
            doc.add_heading(f"Answer {i}", 2)
            doc.add_paragraph(item['answer'])
            
            if item.get('citations'):
                doc.add_paragraph(item['citations'], style='Intense Quote')
            
            doc.add_paragraph()  # Spacer
        
        doc.save(str(filepath))
        logger.info(f"Chat history DOCX exported: {filepath}")
        return str(filepath)


# Example usage
if __name__ == "__main__":
    exporter = ExportHandler()
    
    # Test export
    test_query = "What is distillation?"
    test_answer = "Distillation is a separation process..."
    test_citations = "1. **Chemical Engineering Handbook** (Page 45)\n2. **Unit Operations** (Page 123)"
    
    pdf_path = exporter.export_to_pdf(test_query, test_answer, test_citations, 'book_based')
    print(f"Test PDF: {pdf_path}")
